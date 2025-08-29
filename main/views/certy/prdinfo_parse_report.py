import re
from docx import Document

# ── 유틸 (agreement 모듈과 동일 내용) ─────────────────
def normalize_spaces(s: str) -> str:
    if s is None:
        return ""
    return re.sub(r"\s+", " ", str(s).replace("\u200b", " ").replace("\xa0", " ")).strip()

def flat(s: str) -> str:
    return re.sub(r"\s+", "", normalize_spaces(s))

def text_lines_from_docx(doc: Document):
    lines = []
    for p in doc.paragraphs:
        t = normalize_spaces(p.text)
        if t:
            lines.append(t)
    for tbl in doc.tables:
        for row in tbl.rows:
            row_texts = [normalize_spaces(c.text) for c in row.cells]
            for t in row_texts:
                if t:
                    for sub in [normalize_spaces(x) for x in re.split(r"[\r\n]+", t) if normalize_spaces(x)]:
                        lines.append(sub)
            for i in range(0, len(row_texts)-1):
                l, v = row_texts[i], row_texts[i+1]
                if l or v:
                    lines.append(f"{l} : {v}".strip(" :"))
    return lines

# ── 날짜 패턴 ────────────────────────────
_DATE_PATTERNS = [
    re.compile(r"\d{4}\s*년\s*\d{1,2}\s*월\s*\d{1,2}\s*일"),
    re.compile(r"\d{4}\.\d{1,2}\.\d{1,2}"),
    re.compile(r"\b\d{1,2}/\d{1,2}\b"),
]

# ── 메인: 성적서 파서 ─────────────────────
def extract_process2_docx_overview(byts, filename):
    doc = Document(byts)
    lines = text_lines_from_docx(doc)
    if not lines:
        return {"시험기간": [], "개요 및 특성(설명)": "", "개요 및 특성(주요 기능)": [], "소요일수": 0, "파일명": filename}

    # "6. 시험기간" 블록 추출
    start_idx, end_idx = 0, len(lines)
    for i, t in enumerate(lines):
        if re.search(r"6[.)]?\s*시험기간", t):
            start_idx = i
            break
    for i, t in enumerate(lines[start_idx:], start=start_idx):
        if re.search(r"7[.)]?\s*시험방법", t):
            end_idx = i
            break

    block_text = "\n".join(lines[start_idx:end_idx])  # 줄바꿈 유지
    period_pat = re.compile(
        r"(?:\((?P<label>[^)]+)\)\s*)?"
        r"(?P<start>\d{4}[^~\n]+?(?:일|[0-9]))"
        r"\s*[~\-]?\s*[\r\n ]+"
        r"(?P<end>\d{4}[^)\n]+?(?:일|[0-9]))"
    )
    exam_periods = []
    for m in period_pat.finditer(block_text):
        label = normalize_spaces(m.group("label") or "")
        s, e = normalize_spaces(m.group("start")), normalize_spaces(m.group("end"))
        line = f"({label}) {s} ~ {e}" if label else f"{s} ~ {e}"
        if line not in exam_periods:
            exam_periods.append(line)

    # 폴백: 날짜 포함 줄
    if not exam_periods:
        cap_lines = []
        for t in lines[start_idx:end_idx]:
            if any(p.search(t) for p in _DATE_PATTERNS):
                cap_lines.append(normalize_spaces(t))
        exam_periods = list(dict.fromkeys(cap_lines))

    # 개요/주요 기능
    whole = " ".join(lines)
    desc = ""
    m = re.search(r"본\s*제품은(.*?)(?:으로\s*주요\s*기능은\s*다음과\s*같다)", whole)
    if m:
        desc = normalize_spaces(m.group(1))

    feats = []
    m2 = re.search(r"다음과\s*같다[^\S\r\n]*[:：]?(.*?)(?:※\s*상세기능은|$)", whole, re.S)
    if m2:
        block = normalize_spaces(m2.group(1))
        for seg in re.split(r"[•\-\∙\·\u2219;]|[\r\n]+", block):
            s = normalize_spaces(seg)
            if s:
                feats.append(s)

    # 소요일수(표 합)
    total_days = 0
    for tbl in doc.tables:
        header_idx, days_col = None, None
        for r, row in enumerate(tbl.rows):
            cells = [normalize_spaces(c.text) for c in row.cells]
            if any("소요일수" in x for x in cells):
                header_idx = r
                for c_idx, txt in enumerate(cells):
                    if "소요일수" in txt:
                        days_col = c_idx
                        break
                break
        if header_idx is not None and days_col is not None:
            for row in tbl.rows[header_idx+1:]:
                t = normalize_spaces(row.cells[days_col].text)
                t = re.sub(r"[^\d]", "", t)
                if t.isdigit():
                    total_days += int(t)

    return {
        "시험기간": exam_periods,
        "개요 및 특성(설명)": desc,
        "개요 및 특성(주요 기능)": feats,
        "소요일수": total_days,
        "파일명": filename,
    }
