import re
from docx import Document

# ── 유틸 ─────────────────────────────────
def normalize_spaces(s: str) -> str:
    if s is None:
        return ""
    return re.sub(r"\s+", " ", str(s).replace("\u200b", " ").replace("\xa0", " ")).strip()

def flat(s: str) -> str:
    return re.sub(r"\s+", "", normalize_spaces(s))

def text_lines_from_docx(doc: Document):
    lines = []
    for p in doc.paragraphs:
        t = normalize_spaces(p.text)
        if t:
            lines.append(t)
    for tbl in doc.tables:
        for row in tbl.rows:
            row_texts = [normalize_spaces(c.text) for c in row.cells]
            for t in row_texts:
                if t:
                    for sub in [normalize_spaces(x) for x in re.split(r"[\r\n]+", t) if normalize_spaces(x)]:
                        lines.append(sub)
            for i in range(0, len(row_texts)-1):
                l, v = row_texts[i], row_texts[i+1]
                if l or v:
                    lines.append(f"{l} : {v}".strip(" :"))
    return lines

def find_label_value_in_tables(doc: Document, label_patterns):
    found = {}
    for tbl in doc.tables:
        for row in tbl.rows:
            texts = [normalize_spaces(c.text) for c in row.cells]
            # 한 셀 내 '라벨: 값'
            for t in texts:
                for key, creg in label_patterns.items():
                    if creg.search(t):
                        m = re.split(r"[:：]\s*", t, maxsplit=1)
                        if len(m) == 2 and normalize_spaces(m[1]):
                            found.setdefault(key, normalize_spaces(m[1]))
            # 인접 셀
            for i in range(0, len(texts)-1):
                l, v = texts[i], texts[i+1]
                for key, creg in label_patterns.items():
                    if creg.search(l) and normalize_spaces(v):
                        found.setdefault(key, normalize_spaces(v))
    return found

# ── 라벨 패턴 ────────────────────────────
_1_LABELS_REGEX = {
    "국문명": re.compile(r"\b국문명\b", re.I),
    "영문명": re.compile(r"\b영문명\b", re.I),
    "국문명:": re.compile(r"\b국문명\s*:\b"),
    "영문명:": re.compile(r"\b영문명\s*:\b"),
    "사업자등록번호": re.compile(r"사업자\s*등록\s*번호", re.I),
    "법인등록번호": re.compile(r"법인\s*등록\s*번호", re.I),
    "대표자": re.compile(r"\b대표자\b"),
    "대표자 E- Mail": re.compile(r"대표자\s*E-?\s*Mail", re.I),
    "대표자 E-Mail": re.compile(r"대표자\s*E-?\s*Mail", re.I),
    "대표 전화번호": re.compile(r"대표\s*전화\s*번호", re.I),
    "홈페이지": re.compile(r"\b홈페이지\b", re.I),
    "주 소": re.compile(r"주\s*소", re.I),
    "담당자-성 명": re.compile(r"담당자.*성.?명"),
    "담당자-전화번호": re.compile(r"담당자.*전화"),
    "담당자-Mobile": re.compile(r"담당자.*(Mobile|휴대전화|핸드폰)"),
    "담당자-E- Mail": re.compile(r"담당자.*Mail|담당자.*이메일", re.I),
    "담당자-FAX번호": re.compile(r"담당자.*FAX|담당자.*팩스", re.I),
    "담당자-부서/직급": re.compile(r"담당자.*부서.*|담당자.*직급"),
    "제조자": re.compile(r"\b제조자\b"),
    "제조국가": re.compile(r"\b제조\s*국가\b"),
    "시험신청번호": re.compile(r"\b시험\s*신청\s*번호\b"),
}

def _empty_process1():
    return {
        "시험신청번호": "",
        "성적서 구분": "",
        "국문명": "",
        "영문명": "",
        "사업자등록번호": "",
        "법인등록번호": "",
        "대표자": "",
        "대표자 E-Mail": "",
        "대표 전화번호": "",
        "홈페이지": "",
        "주 소": "",
        "담당자-성 명": "",
        "담당자-전화번호": "",
        "담당자-Mobile": "",
        "담당자-E- Mail": "",
        "담당자-FAX번호": "",
        "담당자-부서/직급": "",
        "국문명:": "",
        "영문명:": "",
        "제조자": "",
        "제조국가": "",
    }

# ── 메인: 합의서 파서 ─────────────────────
def extract_process1_docx_basic(byts, filename):
    doc = Document(byts)
    raw_lines = text_lines_from_docx(doc)
    if not raw_lines:
        return _empty_process1()

    found = find_label_value_in_tables(doc, _1_LABELS_REGEX)

    # 본문 보조 추출: "라벨: 값" 또는 "라벨\n값"
    for line in raw_lines:
        t = normalize_spaces(line)
        for key, creg in _1_LABELS_REGEX.items():
            if creg.search(t):
                m = re.split(r"[:：]\s*", t, maxsplit=1)
                if len(m) == 2 and normalize_spaces(m[1]):
                    found.setdefault(key, normalize_spaces(m[1]))
                else:
                    parts = t.split()
                    if len(parts) >= 2 and not found.get(key):
                        found[key] = normalize_spaces(parts[-1])

    # 시험신청번호 보정
    if "시험신청번호" not in found or not found.get("시험신청번호"):
        for line in raw_lines:
            m = re.search(r"(시험\s*신청\s*번호)\s*[:：]?\s*([A-Za-z0-9\-\_]+)", line)
            if m:
                found["시험신청번호"] = normalize_spaces(m.group(2))
                break

    # 성적서 구분(체크마크)
    score_kind = ""
    MARK = r"(V|√|✔|✓)"
    for line in raw_lines:
        L = normalize_spaces(line)
        if re.search(r"TTA\s*성적서", L, re.I) and re.search(MARK, L):
            score_kind = "TTA 성적서"
        if re.search(r"KOLAS\s*성적서", L, re.I) and re.search(MARK, L):
            score_kind = "KOLAS 성적서"

    out = _empty_process1()
    out.update({
        "시험신청번호": found.get("시험신청번호", ""),
        "성적서 구분": score_kind,
        "국문명": found.get("국문명", ""),
        "영문명": found.get("영문명", ""),
        "사업자등록번호": found.get("사업자등록번호", ""),
        "법인등록번호": found.get("법인등록번호", ""),
        "대표자": found.get("대표자", ""),
        "대표자 E-Mail": found.get("대표자 E- Mail", "") or found.get("대표자 E-Mail", ""),
        "대표 전화번호": found.get("대표 전화번호", ""),
        "홈페이지": found.get("홈페이지", ""),
        "주 소": found.get("주 소", ""),
        "담당자-성 명": found.get("담당자-성 명", ""),
        "담당자-전화번호": found.get("담당자-전화번호", ""),
        "담당자-Mobile": found.get("담당자-Mobile", ""),
        "담당자-E- Mail": found.get("담당자-E- Mail", "") or found.get("담당자-E-Mail", ""),
        "담당자-FAX번호": found.get("담당자-FAX번호", ""),
        "담당자-부서/직급": found.get("담당자-부서/직급", ""),
        "국문명:": found.get("국문명:", ""),
        "영문명:": found.get("영문명:", ""),
        "제조자": found.get("제조자", ""),
        "제조국가": found.get("제조국가", ""),
    })
    return out
