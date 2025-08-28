# -*- coding: utf-8 -*-
from __future__ import annotations
import os, re, json, tempfile
from pathlib import Path
from typing import Dict, Any
from django.http import JsonResponse, HttpResponseBadRequest
from django.views.decorators.http import require_POST
from django.conf import settings

from docx import Document
from PyPDF2 import PdfReader
import openpyxl
from openpyxl.utils import get_column_letter

# Luckysheet로 불러오는 원본 엑셀(템플릿)
ORIGIN_XLSX_PATH = Path(getattr(settings, "BASE_DIR")) / "main/data/prdinfo.xlsx"

# -------------------- 파일 파서 --------------------
def extract_text(fp: Path) -> str:
    fp = Path(fp)
    ext = fp.suffix.lower()
    if ext == ".pdf":
        return "\n".join(p.extract_text() or "" for p in PdfReader(str(fp)).pages)
    if ext == ".docx":
        doc = Document(str(fp))
        parts = [p.text.strip() for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text.strip())
        return "\n".join(p for p in parts if p)
    raise ValueError(f"지원하지 않는 형식: {fp.name} ({ext})")

def extract_defect_counts(xlsx: Path) -> Dict[str, int]:
    wb = openpyxl.load_workbook(str(xlsx), data_only=True)
    keys_q = ["기능적합성","성능효율성","호환성","사용성","신뢰성","보안성","유지보수성","이식성","일반적 요구사항"]
    keys_d = ["H","M","L"]
    result = {k: 0 for k in keys_q + keys_d}
    for sh in wb.worksheets:
        header = {(c.value or "").strip(): i+1 for i, c in enumerate(sh[1])}
        q_col = header.get("품질특성"); d_col = header.get("결함정도")
        if not q_col or not d_col: continue
        for r in range(2, sh.max_row+1):
            qv = sh.cell(r, q_col).value
            if isinstance(qv, str):
                qv = qv.strip()
                if qv in result: result[qv]+=1
                if "일반적" in qv: result["일반적 요구사항"] += 1
            dv = sh.cell(r, d_col).value
            if isinstance(dv, str) and dv in keys_d: result[dv]+=1
    return result

def extract_product_description_and_features(text: str) -> tuple[str, str]:
    desc = ""; features = []
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    for line in lines:
        if not desc and "본 제품은" in line and "주요 기능은" in line:
            m = re.search(r"본\s*제품은\s*(.*?)\s*주요\s*기능은\s*다음과 같다", line)
            if m: desc = m.group(1).strip()
        elif line.startswith("-") and "※" not in line:
            features.append(line.strip())
        elif "※" in line:
            break
    return desc, "\n".join(features)

# -------------------- 템플릿 → 채우기 맵 --------------------
def _addr(r: int, c: int) -> str: return f"{get_column_letter(c)}{r}"

def build_fill_map_from_template(data: Dict[str, Any], tpl_path: Path) -> Dict[str, Dict[str, Any]]:
    wb = openpyxl.load_workbook(str(tpl_path), data_only=True)
    fill_map: Dict[str, Dict[str, Any]] = {}
    for sh in wb.worksheets:
        m = fill_map.setdefault(sh.title, {})

        if sh.title == "제품 정보 요청":
            placeholders = {
                "㈜000\nAAA Co., Ltd.": data.get("회사명",""),
                "000 v3.5\nAAA v3.5":    data.get("제품명",""),
                "GS-A-00-000":           data.get("GS번호",""),
                "계약된 총 WD 기재\n(예: 35)\n(메모 참조)": data.get("총WD",""),
                "20XX.XX.XX ~ 20XX.XX.XX": data.get("결과서 기재 시험 기간",""),
                "(000-000)":               data.get("주소",""),
                "㈜000\n/대한민국":        data.get("제조자/제조국가",""),
                "www.000.co.kr":          data.get("홈페이지",""),
            }
            for row in sh.iter_rows():
                for cell in row:
                    if isinstance(cell.value, str):
                        v = cell.value.strip()
                        if v in placeholders and placeholders[v]!="":
                            m[_addr(cell.row, cell.column)] = placeholders[v]

            for row in sh.iter_rows():
                for cell in row:
                    if not isinstance(cell.value, str): continue
                    flat = cell.value.replace("\n","").replace(" ","")
                    if flat.startswith("제품설명"):
                        m[_addr(cell.row+1, cell.column)] = data.get("제품설명(시험 결과서 개요)","")
                    elif flat.startswith("제품주요기능"):
                        m[_addr(cell.row+1, cell.column)] = data.get("제품 주요기능(시험 결과서 개요 부분 주요 기능)","")

            contains_under = ["사업자등록번호","대표자","대표전화","업무담당자","전화번호","팩스"]
            for row in sh.iter_rows():
                for cell in row:
                    if isinstance(cell.value, str):
                        for key in contains_under:
                            if key in cell.value:
                                m[_addr(cell.row+1, cell.column)] = data.get(key,"")

            exact_under = {"이메일":"이메일","대표자이메일":"대표자이메일","성적서 구분":"성적서 구분"}
            for row in sh.iter_rows():
                for cell in row:
                    if isinstance(cell.value, str) and cell.value.strip() in exact_under:
                        k = exact_under[cell.value.strip()]
                        m[_addr(cell.row+1, cell.column)] = data.get(k,"")

        elif sh.title == "결함정보":
            keys = ["기능적합성","성능효율성","호환성","사용성","신뢰성","보안성","유지보수성","이식성","일반적 요구사항","H","M","L"]
            for i, k in enumerate(keys, 3): m[_addr(4,i)] = data.get(k,0)
            for row in sh.iter_rows():
                for cell in row:
                    if isinstance(cell.value, str) and ("결함" in cell.value and "차수" in cell.value):
                        m[_addr(cell.row+2, cell.column)] = data.get("결함차수","1"); break
    return fill_map

# -------------------- 규칙 기반 데이터 생성 --------------------
def build_data(text_score: str, text_agree: str, defect_summary: Dict[str,int],
               score_name: str, defect_filename: str) -> Dict[str, Any]:
    data: Dict[str, Any] = {}
    def find_one(pats, text):
        for p in pats:
            m = re.search(p, text, re.I)
            if m: return m.group(1).strip()
        return ""

    data["회사명"] = (find_one([r"회사명[:：]\s*([^\n]+)", r"상\s*호[:：]\s*([^\n]+)"], text_agree)
                      or find_one([r"회사명[:：]\s*([^\n]+)"], text_score))
    data["제품명"] = (find_one([r"제품명[:：]\s*([^\n]+)"], text_score) or Path(score_name).stem)

    period = find_one([r"(\d{4}\.\d{2}\.\d{2})\s*[~\-]\s*(\d{4}\.\d{2}\.\d{2})"], text_score)
    data["결과서 기재 시험 기간"] = period.replace("-", "~") if period else ""

    for k, pats in {
        "사업자등록번호": [r"사업자등록번호[:：]\s*([0-9\-]+)"],
        "대표자":       [r"대표자[:：]\s*([^\n]+)"],
        "대표전화":     [r"대표전화[:：]\s*([^\n]+)"],
        "업무담당자":   [r"업무담당자[:：]\s*([^\n]+)"],
        "전화번호":     [r"전화번호[:：]\s*([^\n]+)"],
        "팩스":         [r"팩스[:：]\s*([^\n]+)"],
        "이메일":       [r"이메일[:：]\s*([^\n]+)"],
        "대표자이메일": [r"대표자이메일[:：]\s*([^\n]+)"],
        "주소":         [r"주소[:：]\s*([^\n]+)"],
        "제조자/제조국가":[r"(제조자\/제조국가)[:：]\s*([^\n]+)"],
        "홈페이지":     [r"(홈페이지|웹사이트)[:：]\s*([^\n]+)"],
        "성적서 구분":  [r"(성적서\s*구분)[:：]\s*([^\n]+)"],
        "총WD":         [r"총\s*WD[:：]?\s*(\d+)", r"WD\s*합계[:：]?\s*(\d+)"],
    }.items():
        val = find_one(pats, text_score + "\n" + text_agree)
        data[k] = val or ""

    desc, feats = extract_product_description_and_features(text_score)
    data.setdefault("제품설명(시험 결과서 개요)", desc)
    data.setdefault("제품 주요기능(시험 결과서 개요 부분 주요 기능)", feats)

    m = re.search(r"(GS-[A-Z]-\d{2}-\d{4})", score_name)
    data["GS번호"] = m.group(1) if m else "GS-A-00-0000"

    data.update(defect_summary)
    ver = re.search(r"v(\d+)", Path(defect_filename).stem, re.I)
    data["결함차수"] = ver.group(1) if ver else "1"
    return data

# -------------------- Django View (AJAX JSON 전용) --------------------
@require_POST
def generate_prdinfo(request):
    # 로그인 리다이렉트(302) 방지: JSON 401 반환
    #if not request.user.is_authenticated:
    #    return JsonResponse({'detail': 'Unauthorized'}, status=401)

    files = request.FILES.getlist("file")
    if not files or len(files) != 3:
        return HttpResponseBadRequest("파일 3개를 업로드하세요.")

    ALLOWED = {".pdf", ".docx", ".xlsx"}
    kw = {"성적서": None, "합의서": None, "결함": None}
    tmpdir = Path(tempfile.mkdtemp(prefix="prdinfo_"))

    for f in files:
        name = f.name; ext = Path(name).suffix.lower()
        if ext not in ALLOWED:
            return HttpResponseBadRequest("허용 확장자는 pdf, docx, xlsx 입니다.")
        p = tmpdir / name
        with p.open("wb") as out:
            for chunk in f.chunks(): out.write(chunk)
        if "성적서" in name: kw["성적서"] = p
        elif "합의서" in name: kw["합의서"] = p
        elif ("결함리포트" in name) or ("결함" in name): kw["결함"] = p

    if not all(kw.values()):
        return HttpResponseBadRequest("합의서, 성적서, 결함리포트를 모두 업로드했는지 확인하세요.")

    # 1) 파싱
    text_score = extract_text(kw["성적서"])
    text_agree = extract_text(kw["합의서"])
    defect_summary = extract_defect_counts(kw["결함"])

    # 2) 규칙 기반 데이터
    data = build_data(
        text_score=text_score,
        text_agree=text_agree,
        defect_summary=defect_summary,
        score_name=kw["성적서"].name,
        defect_filename=kw["결함"].name,
    )

    # 3) 템플릿 → 좌표맵
    tpl_path = ORIGIN_XLSX_PATH if ORIGIN_XLSX_PATH.exists() else \
               list(Path.cwd().glob("GS-*_제품_정보_요청_첨부_v*.xlsx"))[0]
    fill_map = build_fill_map_from_template(data, tpl_path)
    return JsonResponse({"fillMap": fill_map, "gsNumber": data.get("GS번호","")})
