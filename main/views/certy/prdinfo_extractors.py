# certy/parsers/prdinfo_extractors.py
"""
프롬포트 스펙을 충실히 반영한 파서 모음.
- 1번 과정(.docx, '합의서'): 기업/제품 기본정보 추출
- 2번 과정(.docx, '성적서'): 시험기간/개요/주요기능/소요일수
- 3번 과정(.xlsx, '결함리포트'): 결함 집계(품질특성별/결함정도별)
주의:
- '오류' 문구는 뷰에서만 사용하며, 파서는 가능한 한 결과를 빈문자열/0으로 채움.
- 3번은 시트/블록을 못 찾아도 0 보정하여 동일 구조 반환(파일 로딩 실패만 예외).  # 스펙
"""

import re
from collections import defaultdict
from docx import Document
from openpyxl import load_workbook

# ─────────────────────────────────────────────────────
# 공통 유틸
# ─────────────────────────────────────────────────────

def normalize_spaces(s: str) -> str:
    if s is None:
        return ""
    # ZWSP/NBSP/탭/다중 공백 → 한 칸, 앞뒤 trim
    return re.sub(r"\s+", " ", str(s).replace("\u200b", " ").replace("\xa0", " ")).strip()

def flat(s: str) -> str:
    # 라벨 비교용: 공백류 제거
    return re.sub(r"\s+", "", normalize_spaces(s))

def text_lines_from_docx(doc: Document):
    """문서의 모든 문장/표 셀을 줄 단위 배열로 평탄화."""
    lines = []
    for p in doc.paragraphs:
        t = normalize_spaces(p.text)
        if t:
            lines.append(t)
    for tbl in doc.tables:
        for row in tbl.rows:
            row_texts = [normalize_spaces(c.text) for c in row.cells]
            # 각 셀의 전체 문장들도 라인으로 넣고,
            for i, t in enumerate(row_texts):
                if t:
                    for sub in [normalize_spaces(x) for x in re.split(r"[\r\n]+", t) if normalize_spaces(x)]:
                        lines.append(sub)
            # 라벨-값 페어 추정 위해 "(cell[i], cell[i+1])"도 라인화
            for i in range(0, len(row_texts)-1):
                l, v = row_texts[i], row_texts[i+1]
                if l or v:
                    lines.append(f"{l} : {v}".strip(" :"))
    return lines

def find_label_value_in_tables(doc: Document, label_patterns):
    """
    표 기준의 라벨 → 우측 값 탐색(셀 병합/한 셀 내 '라벨: 값' 포함).
    label_patterns: {key: compiled_regex_for_label}
    반환: {key: value}
    """
    found = {}
    for tbl in doc.tables:
        # 행 단위로 좌/우 칸 보며 매칭
        for row in tbl.rows:
            texts = [normalize_spaces(c.text) for c in row.cells]
            # 한 셀 내 '라벨: 값' 케이스
            for t in texts:
                for key, creg in label_patterns.items():
                    if creg.search(t):
                        # "라벨: 값" 형태면 콜론 뒤를 우선
                        m = re.split(r"[:：]\s*", t, maxsplit=1)
                        if len(m) == 2 and normalize_spaces(m[1]):
                            found.setdefault(key, normalize_spaces(m[1]))
            # 인접 우측 셀 케이스
            for i in range(0, len(texts)-1):
                l, v = texts[i], texts[i+1]
                for key, creg in label_patterns.items():
                    if creg.search(l) and normalize_spaces(v):
                        found.setdefault(key, normalize_spaces(v))
    return found

# ─────────────────────────────────────────────────────
# 1번 과정: .docx 기본정보 추출
# ─────────────────────────────────────────────────────

_1_LABELS_REGEX = {
    # 기업/기관명 (콜론 없는 라벨)
    "국문명": re.compile(r"\b국문명\b", re.I),
    "영문명": re.compile(r"\b영문명\b", re.I),
    # 제품명 (콜론 있는 라벨)
    "국문명:": re.compile(r"\b국문명\s*:\b"),
    "영문명:": re.compile(r"\b영문명\s*:\b"),
    "사업자등록번호": re.compile(r"사업자\s*등록\s*번호", re.I),
    "법인등록번호": re.compile(r"법인\s*등록\s*번호", re.I),
    "대표자": re.compile(r"\b대표자\b"),
    "대표자 E- Mail": re.compile(r"대표자\s*E-?\s*Mail", re.I),
    "대표자 E-Mail": re.compile(r"대표자\s*E-?\s*Mail", re.I),  # 형태 보정
    "대표 전화번호": re.compile(r"대표\s*전화\s*번호", re.I),
    "홈페이지": re.compile(r"\b홈페이지\b", re.I),
    "주 소": re.compile(r"주\s*소", re.I),
    "담당자-성 명": re.compile(r"담당자\s*-\s*성\s*명|담당자\s*성명", re.I),
    "담당자-전화번호": re.compile(r"담당자\s*-\s*전화\s*번호|담당자\s*전화번호", re.I),
    "담당자-Mobile": re.compile(r"담당자\s*-\s*Mobile|담당자\s*휴대전화|담당자\s*핸드폰", re.I),
    "담당자-E- Mail": re.compile(r"담당자\s*-\s*E-?\s*Mail|담당자\s*이메일", re.I),
    "담당자-FAX번호": re.compile(r"담당자\s*-\s*FAX\s*번호|담당자\s*팩스", re.I),
    "담당자-부서/직급": re.compile(r"담당자\s*-\s*부서\s*/\s*직급|담당자\s*부서\s*/\s*직급|담당자\s*부서|담당자\s*직급", re.I),

    # 추가 필드
    "제조자": re.compile(r"\b제조자\b"),
    "제조국가": re.compile(r"\b제조\s*국가\b"),
    "시험신청번호": re.compile(r"\b시험\s*신청\s*번호\b"),
}

def extract_process1_docx_basic(byts, filename):
    """
    합의서 .docx 에서 기본정보 추출.
    스펙:
    - '성적서 구분' = 같은 행/줄에서 'TTA 성적서' 또는 'KOLAS 성적서'와 체크('V','√','✔','✓')가 함께 보이면 해당 값
    - '국문명/영문명' vs '국문명:/영문명:' 구분 유지
    - 문서가 비어 있으면 모든 필드 "" 반환
    - 파싱 실패는 뷰에서 오류 문자열 처리
    """
    doc = Document(byts)
    # 실질 텍스트 유무 판단
    raw_lines = text_lines_from_docx(doc)
    if not raw_lines:
        return _empty_process1()

    # 표 기반 라벨 추출
    found = find_label_value_in_tables(doc, _1_LABELS_REGEX)

    # 표 밖 본문에서도 보조 추출(예: '대표자 E-Mail : xxx@yyy', '시험신청번호: GS-...')
    for line in raw_lines:
        t = normalize_spaces(line)
        for key, creg in _1_LABELS_REGEX.items():
            if creg.search(t):
                m = re.split(r"[:：]\s*", t, maxsplit=1)
                if len(m) == 2 and normalize_spaces(m[1]):
                    found.setdefault(key, normalize_spaces(m[1]))
    # 시험신청번호 별도 보정(본문에만 있을 수 있음)
    if "시험신청번호" not in found:
        for line in raw_lines:
            m = re.search(r"(시험\s*신청\s*번호)\s*[:：]?\s*([A-Za-z0-9\-\_]+)", line)
            if m:
                found["시험신청번호"] = normalize_spaces(m.group(2))
                break

    # 성적서 구분 탐지(체크마크 포함 줄)
    score_kind = ""
    MARK = r"(V|√|✔|✓)"
    for line in raw_lines:
        L = normalize_spaces(line)
        if re.search(r"TTA\s*성적서", L, re.I) and re.search(MARK, L):
            score_kind = "TTA 성적서"
        if re.search(r"KOLAS\s*성적서", L, re.I) and re.search(MARK, L):
            score_kind = "KOLAS 성적서"
    # 모호하면 빈 문자열 유지  # 스펙
    # 기타 필드 구성
    out = _empty_process1()
    out.update({
        "시험신청번호": found.get("시험신청번호", ""),
        "성적서 구분": score_kind,
        "국문명": found.get("국문명", ""),
        "영문명": found("영문명", ""),
        "사업자등록번호": found.get("사업자등록번호", ""),
        "법인등록번호": found.get("법인등록번호", ""),
        "대표자": found.get("대표자", ""),
        "대표자 E-Mail": found.get("대표자 E- Mail", "") or found.get("대표자 E-Mail", ""),
        "대표 전화번호": found.get("대표 전화번호", ""),
        "홈페이지": found.get("홈페이지", ""),
        "주 소": found.get("주 소", ""),
        "담당자-성 명": found.get("담당자-성 명", ""),
        "담당자-전화번호": found.get("담당자-전화번호", ""),
        "담당자-Mobile": found.get("담당자-Mobile", ""),
        "담당자-E- Mail": found.get("담당자-E- Mail", ""),
        "담당자-FAX번호": found.get("담당자-FAX번호", ""),
        "담당자-부서/직급": found.get("담당자-부서/직급", ""),
        "국문명:": found.get("국문명:", ""),
        "영문명:": found.get("영문명:", ""),
        "제조자": found.get("제조자", ""),
        "제조국가": found.get("제조국가", ""),
    })
    return out

def _empty_process1():
    # 스펙에 명시된 키(제품·기업 구분 포함) - 누락 허용 없이 항상 존재
    return {
        "시험신청번호": "",
        "성적서 구분": "",
        "국문명": "",
        "영문명": "",
        "사업자등록번호": "",
        "법인등록번호": "",
        "대표자": "",
        "대표자 E-Mail": "",
        "대표 전화번호": "",
        "홈페이지": "",
        "주 소": "",
        "담당자-성 명": "",
        "담당자-전화번호": "",
        "담당자-Mobile": "",
        "담당자-E- Mail": "",
        "담당자-FAX번호": "",
        "담당자-부서/직급": "",
        "국문명:": "",
        "영문명:": "",
        "제조자": "",
        "제조국가": "",
    }

# ─────────────────────────────────────────────────────
# 2번 과정: .docx 개요/기간/소요일수
# ─────────────────────────────────────────────────────

_DATE_PATTERNS = [
    re.compile(r"\d{4}\s*년\s*\d{1,2}\s*월\s*\d{1,2}\s*일"),
    re.compile(r"\d{4}\.\d{1,2}\.\d{1,2}"),
    re.compile(r"\b\d{1,2}/\d{1,2}\b"),
]

def extract_process2_docx_overview(byts, filename):
    doc = Document(byts)
    lines = text_lines_from_docx(doc)
    if not lines:
        return {
            "시험기간": [],
            "개요 및 특성(설명)": "",
            "개요 및 특성(주요 기능)": [],
            "소요일수": 0,
            "파일명": filename,
        }

    # 2-1) "7. 시험방법" 처음 등장 전까지의 날짜 포함 라인 수집
    cap_lines = []
    upper_idx = len(lines)
    for i, t in enumerate(lines):
        if "7. 시험방법" in t:
            upper_idx = i
            break
    for t in lines[:upper_idx]:
        if any(p.search(t) for p in _DATE_PATTERNS):
            cap_lines.append(normalize_spaces(t))
    # 중복 제거
    seen = set(); exam_periods = []
    for t in cap_lines:
        if t not in seen:
            exam_periods.append(t); seen.add(t)

    # 2-2) 개요·특성(설명): "본 제품은" 이후 ~ "으로 주요 기능은 다음과 같다" 직전
    whole = " ".join(lines)
    desc = ""
    m = re.search(r"본\s*제품은(.*?)(?:으로\s*주요\s*기능은\s*다음과\s*같다)", whole)
    if m:
        desc = normalize_spaces(m.group(1))

    # 2-3) 주요 기능: "다음과 같다" 이후 ~ "※ 상세기능은" 직전, 불릿/개행으로 분해
    feats = []
    m2 = re.search(r"다음과\s*같다[^\S\r\n]*[:：]?(.*?)(?:※\s*상세기능은|$)", whole, re.S)
    if m2:
        block = normalize_spaces(m2.group(1))
        # 불릿/줄바꿈/세미콜론 등으로 항목화
        for seg in re.split(r"[•\-\∙\·\u2219;]|[\r\n]+", block):
            s = normalize_spaces(seg)
            if s:
                feats.append(s)

    # 2-4) 소요일수: 모든 표에서 "소요일수" 헤더 열의 숫자 합
    total_days = 0
    for tbl in doc.tables:
        # 헤더 행 탐색
        header_idx = None
        days_col = None
        for r, row in enumerate(tbl.rows):
            cells = [normalize_spaces(c.text) for c in row.cells]
            flat_cells = [flat(x) for x in cells]
            if any("소요일수" in x for x in cells):
                header_idx = r
                # 동적 인덱스(열 순서가 바뀌어도 동작)
                for c_idx, txt in enumerate(cells):
                    if "소요일수" in txt:
                        days_col = c_idx
                        break
                break
        if header_idx is not None and days_col is not None:
            # 숫자 행 합
            for row in tbl.rows[header_idx+1:]:
                t = normalize_spaces(row.cells[days_col].text)
                t = re.sub(r"[^\d]", "", t)  # 콤마/문자 제거
                if t.isdigit():
                    total_days += int(t)

    return {
        "시험기간": exam_periods,
        "개요 및 특성(설명)": desc,
        "개요 및 특성(주요 기능)": feats,
        "소요일수": total_days,
        "파일명": filename,
    }

# ─────────────────────────────────────────────────────
# 3번 과정: .xlsx 결함 집계
# ─────────────────────────────────────────────────────

# 라벨 정규화 매핑(스펙)
QUAL_LABEL_MAP = {
    "적합성": ["기능적합성"],
    "효율성": ["성능효율성"],
    "호환성": ["호환성"],
    "사용성": ["사용성"],
    "신뢰성": ["신뢰성"],
    "보안성": ["보안성"],
    "유지보수성": ["유지보수성"],
    "이식성": ["이식성"],
    "요구사항": ["일반적요구사항"],
}
DEG_LABELS = ["High", "Medium", "Low"]  # 결함정도

def extract_process3_xlsx_defects(byts, filename):
    """
    '시험분석자료' 시트에서:
      - 블록 A: 품질특성별 결함내역 → 각 항목 '수정전/최종'
      - 블록 B: 결함정도별 결함내역 → High/Medium/Low '수정전/최종'
    헤더(수정전/최종) 위치는 동적으로 탐지. 미발견은 0 보정.
    '결함차수'는 파일명 내 'vX.Y' → X 정수로 추출(없으면 0).
    """
    wb = load_workbook(byts, data_only=True)
    sh = wb["시험분석자료"] if "시험분석자료" in wb.sheetnames else None

    # 기본 뼈대(0 보정)
    out = {
        "결함차수": _defect_round_from_filename(filename),
        "적합성": {"수정전": 0, "최종": 0},
        "효율성": {"수정전": 0, "최종": 0},
        "호환성": {"수정전": 0, "최종": 0},
        "사용성": {"수정전": 0, "최종": 0},
        "신뢰성": {"수정전": 0, "최종": 0},
        "보안성": {"수정전": 0, "최종": 0},
        "유지보수성": {"수정전": 0, "최종": 0},
        "이식성": {"수정전": 0, "최종": 0},
        "요구사항": {"수정전": 0, "최종": 0},
        "High": {"수정전": 0, "최종": 0},
        "Medium": {"수정전": 0, "최종": 0},
        "Low": {"수정전": 0, "최종": 0},
    }
    if sh is None:
        return out  # 스펙: 시트 없어도 0 보정으로 반환

    # 시트 전체 텍스트 스캔해서 블록 추정(완화 매칭)
    values = [[(cell.value if cell.value is not None else "") for cell in row] for row in sh.iter_rows()]
    H, W = len(values), max((len(r) for r in values), default=0)

    def txt(r, c):
        try:
            return normalize_spaces(values[r][c])
        except Exception:
            return ""

    def flat_txt(r, c):
        return flat(txt(r, c))

    # 헤더 찾기: "수정전", "최종"
    def find_header_cols(start_r):
        for r in range(start_r, min(start_r+5, H)):
            row = [txt(r, c) for c in range(W)]
            try:
                pre_idx = next(i for i, v in enumerate(row) if "수정전" in v)
                fin_idx = next(i for i, v in enumerate(row) if "최종" in v)
                return r, pre_idx, fin_idx
            except StopIteration:
                continue
        return None

    # 블록 A: "품질특성별 결함내역" 시작 행
    startA = None
    for r in range(H):
        line = "".join(txt(r, c) for c in range(W))
        if "품질" in line and "특성" in line and "결함" in line:
            startA = r
            break
    if startA is not None:
        hdr = find_header_cols(startA+1) or find_header_cols(startA+2) or find_header_cols(startA)
        if hdr:
            hr, preC, finC = hdr
            # 항목 행 스캔 ('계' 또는 빈 라벨에서 종료)
            for r in range(hr+1, H):
                label = txt(r, 0)  # 1열에 라벨이 온다고 가정(일반적)
                if not label or "계" in label:
                    break
                # 라벨 정규화 및 매핑
                norm = flat(label)
                for key, aliases in QUAL_LABEL_MAP.items():
                    if flat(key) in norm or any(a and flat(a) in norm for a in aliases):
                        out[key] = {
                            "수정전": _to_int(txt(r, preC)),
                            "최종": _to_int(txt(r, finC)),
                        }
                        break

    # 블록 B: "결함정도별 결함내역"
    startB = None
    for r in range(H):
        line = "".join(txt(r, c) for c in range(W))
        if "결함정도" in line and "결함내역" in line:
            startB = r
            break
    if startB is not None:
        hdr = find_header_cols(startB+1) or find_header_cols(startB+2) or find_header_cols(startB)
        if hdr:
            hr, preC, finC = hdr
            for r in range(hr+1, H):
                label = txt(r, 0)
                if not label or "계" in label:
                    break
                for key in DEG_LABELS:
                    if key.lower() in label.lower():
                        out[key] = {"수정전": _to_int(txt(r, preC)), "최종": _to_int(txt(r, finC))}
                        break

    return out

def _to_int(s):
    s = normalize_spaces(s)
    if not s:
        return 0
    s = re.sub(r"[^\d\-]", "", s)
    try:
        return int(s)
    except Exception:
        return 0

def _defect_round_from_filename(name: str) -> int:
    """
    파일명 내 'vX.Y' 형태가 있으면 X 정수 반환, 없으면 0.
    """
    m = re.search(r"v(\d+)(?:\.\d+)?", name or "", re.I)
    return int(m.group(1)) if m else 0

# ─────────────────────────────────────────────────────
# fillMap 조립
# ─────────────────────────────────────────────────────

def build_fill_map(obj1: dict, obj2: dict, obj3: dict):
    """
    프론트 매핑 스펙에 따라 Luckysheet 좌표로 변환.
    - 시트 "제품 정보 요청": 1번+2번
    - 시트 "결함정보": 3번
    (한 셀 내 2줄은 '\n')
    """
    # 1번 → "제품 정보 요청"
    b5 = "\n".join([obj1.get("국문명", ""), obj1.get("영문명", "")]).strip("\n")
    h7 = "\n".join([obj1.get("담당자-성 명", ""), obj1.get("담당자-FAX번호", "")]).strip("\n")
    c5 = "\n".join([obj1.get("국문명:", ""), obj1.get("영문명:", "")]).strip("\n")
    dept, title = "", ""
    if obj1.get("담당자-부서/직급"):
        parts = [x.strip() for x in obj1["담당자-부서/직급"].split("/", 1)]
        dept = parts[0] if len(parts) >= 1 else ""
        title = parts[1] if len(parts) >= 2 else ""

    prod_sheet = {
        # 1) 1번 과정 순서 매핑
        "D5": obj1.get("시험신청번호", ""),
        "N5": obj1.get("성적서 구분", ""),
        "B5": b5,                   # (국문명/영문명 2줄)
        "B7": obj1.get("사업자등록번호", ""),
        "C7": obj1.get("법인등록번호", ""),
        "D7": obj1.get("대표자", ""),
        "F7": obj1.get("대표자 E-Mail", ""),
        "E7": obj1.get("대표 전화번호", ""),
        "M7": obj1.get("홈페이지", ""),
        "G7": obj1.get("주 소", ""),
        "H7": h7,                   # (담당자-성 명 / 담당자-FAX번호)
        # 13. 건너뜀
        "I7": obj1.get("담당자-전화번호", ""),
        "K7": obj1.get("담당자-E- Mail", ""),
        "J7": obj1.get("담당자-Mobile", ""),
        "C5": c5,                   # (국문명: / 영문명:)
        "L7": "\n".join([dept, title]).strip("\n"),   # (부서/직급)
        # 2) 2번 과정 매핑
        "K5": "\n".join(obj2.get("시험기간", []) or []),
        "F5": obj2.get("개요 및 특성(설명)", "") or "",
        "G5": "\n".join(obj2.get("개요 및 특성(주요 기능)", []) or []),
        "H5": obj2.get("소요일수", 0),
    }

    # 3번 → "결함정보"
    defect_sheet = {
        "B4": obj3.get("결함차수", 0),
        "C4": obj3.get("적합성", {}).get("수정전", 0),
        "D4": obj3.get("효율성", {}).get("수정전", 0),
        "E4": obj3.get("호환성", {}).get("수정전", 0),
        "F4": obj3.get("사용성", {}).get("수정전", 0),
        "G4": obj3.get("신뢰성", {}).get("수정전", 0),
        "H4": obj3.get("보안성", {}).get("수정전", 0),
        "I4": obj3.get("유지보수성", {}).get("수정전", 0),
        "J4": obj3.get("이식성", {}).get("수정전", 0),
        "K4": obj3.get("요구사항", {}).get("수정전", 0),
        "L4": obj3.get("High", {}).get("수정전", 0),
        "M4": obj3.get("Medium", {}).get("수정전", 0),
        "N4": obj3.get("Low", {}).get("수정전", 0),
    }

    return {
        "제품 정보 요청": prod_sheet,
        "결함정보": defect_sheet,
    }
