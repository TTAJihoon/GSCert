# -*- coding: utf-8 -*-
"""
report_docx_parser.py
- checkreport_api.py 의 build_pages(docx_path, pdf_path=...) 호출과 호환
- DOCX의 문단/표를 추출하고, PDF의 페이지별 텍스트를 기준으로 안정적으로 매핑
- 번호형 헤딩은 label, 나머지 본문은 sen 으로 분류
- 표가 같은 시작셀 텍스트를 가져도 1페이지에 몰리지 않도록 연속표 분산 로직 포함
"""
from __future__ import annotations
import re
import json
from dataclasses import dataclass
from typing import Any, Dict, List, Optional, Tuple, Iterable

# -----------------------
# Utilities / Regex rules
# -----------------------

LABEL_PATTERNS = [
    # 1. 숫자/소숫점 기반 섹션(예: "1. 개요", "2 시험목적", "2.1 시험 대상")
    re.compile(r"^\s*\d+(?:\.\d+)*\s*[).]?\s+"),
    # 2. 각괄호/꺾쇠 첨부(예: "<첨부1>", "(첨부 1)", "<시험환경구성도>")
    re.compile(r"^\s*(?:<\s*첨부[^>]*>|\(첨부[^)]*\)|<[^>]+>)\s*$"),
    # 3. 앞에 숫자와 제목 붙은 목차형 (예: "5. 시험결과11" 같은 ToC 라인)
    re.compile(r"^\s*\d+(?:\.\d+)*\s+.+?\d+\s*$"),
]

def is_label(text: str) -> bool:
    t = (text or "").strip()
    if not t:
        return False
    for pat in LABEL_PATTERNS:
        if pat.match(t):
            return True
    return False

def normalize(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").strip())

def tokenize_for_match(s: str, min_len: int = 2) -> List[str]:
    # 한글/영문/숫자만 추출 후 짧은 토큰 제거
    tokens = re.findall(r"[가-힣A-Za-z0-9]+", s or "")
    return [t for t in tokens if len(t) >= min_len]

# -----------------------
# PDF text extraction
# -----------------------

def extract_pdf_pages_text(pdf_path: str) -> List[str]:
    """
    PyPDF2 로 페이지별 텍스트 추출
    """
    try:
        import PyPDF2
    except Exception as e:
        raise RuntimeError("PyPDF2 미설치: pip install PyPDF2") from e

    reader = PyPDF2.PdfReader(pdf_path)
    pages: List[str] = []
    for p in reader.pages:
        txt = p.extract_text() or ""
        # normalize for matching
        txt = normalize(txt)
        pages.append(txt)
    return pages

# -----------------------
# DOCX extraction
# -----------------------

@dataclass
class DocxParagraph:
    text: str

@dataclass
class DocxTable:
    # cells: [ [row, col, rowspan, colspan, text], ... ]
    cells: List[List[Any]]

DocxItem = Tuple[str, Any]  # ("paragraph", DocxParagraph) | ("table", DocxTable)

def extract_docx_items(docx_path: str) -> List[DocxItem]:
    """
    python-docx 로 문단/표 순서대로 추출
    표 병합정보는 없으면 1,1로 채움
    """
    try:
        from docx import Document
    except Exception as e:
        raise RuntimeError("python-docx 미설치: pip install python-docx") from e

    doc = Document(docx_path)

    # python-docx 는 문단/표 interleaved 순회를 직접 제공하지 않으므로
    # 각 섹션의 _element 하위로 순회하면서 타입 판별
    items: List[DocxItem] = []
    body_elements = []
    for block in doc.element.body.iterchildren():
        body_elements.append(block)

    # paragraph / table 구분
    for el in body_elements:
        if el.tag.endswith("tbl"):  # table
            # 표 변환
            # 간단히 doc.tables 를 다시 돌리면 순서가 섞일 수 있으니 element -> table 객체 재구성
            tbl = None
            for t in doc.tables:
                if t._element is el:
                    tbl = t
                    break
            if not tbl:
                continue
            cells_payload: List[List[Any]] = []
            for r_idx, row in enumerate(tbl.rows, start=1):
                for c_idx, cell in enumerate(row.cells, start=1):
                    text = normalize(cell.text)
                    cells_payload.append([r_idx, c_idx, 1, 1, text])
            items.append(("table", DocxTable(cells=cells_payload)))
        else:
            # paragraph
            # element -> Paragraph 객체 매핑
            para = None
            # doc.paragraphs 를 돌며 같은 element 찾기
            # (성능 이슈가 크지 않으므로 선형탐색 허용)
            for p in doc.paragraphs:
                if p._p is el:
                    para = p
                    break
            if para:
                txt = normalize(para.text)
                if txt:
                    items.append(("paragraph", DocxParagraph(text=txt)))

    return items

# -----------------------
# Matching / Scoring
# -----------------------

def sample_strings_from_table(tbl: DocxTable, max_samples: int = 8) -> List[str]:
    """
    표에서 매칭에 쓸 샘플 문자열 추출.
    - 우선순위: 첫 행/첫 열/머리셀 → 충분치 않으면 본문 셀에서 몇 개 더 채움
    """
    texts: List[str] = []
    # 1) 첫 행
    first_row = min((r for r,_,_,_,_ in tbl.cells), default=1)
    for r,c,rs,cs,tx in tbl.cells:
        if r == first_row and tx:
            texts.append(tx)
    # 2) 첫 열
    first_col = min((c for _,c,_,_,_ in tbl.cells), default=1)
    for r,c,rs,cs,tx in tbl.cells:
        if c == first_col and tx:
            texts.append(tx)
    # 3) 본문 일부 추가
    for r,c,rs,cs,tx in tbl.cells:
        if tx:
            texts.append(tx)
    # 정리
    uniq = []
    seen = set()
    for t in texts:
        t = normalize(t)
        if not t or t in seen:
            continue
        seen.add(t)
        uniq.append(t)
        if len(uniq) >= max_samples:
            break
    return uniq

def score_text_on_page(text: str, page_text: str) -> float:
    """
    간단 토큰 교집합 기반 점수.
    - 텍스트 길이에 따라 최소 토큰수/가중치 보정
    """
    if not text or not page_text:
        return 0.0
    tokens = set(tokenize_for_match(text))
    if not tokens:
        return 0.0
    page_tokens = set(tokenize_for_match(page_text))
    inter = tokens & page_tokens
    # 짧은 문장 과적합 방지: 토큰 길이에 따라 기준 상향
    need = 1 if len(tokens) <= 3 else 2 if len(tokens) <= 8 else 3
    base = 1.0 if len(inter) >= need else 0.0
    # 더 많은 교집합일수록 가중
    return base + 0.2 * max(0, len(inter) - need)

def score_table_on_pages(tbl: DocxTable, pdf_pages: List[str]) -> List[float]:
    samples = sample_strings_from_table(tbl)
    scores: List[float] = []
    for ptxt in pdf_pages:
        s = 0.0
        for i, ss in enumerate(samples):
            # 첫 샘플(헤더일 확률 높음)은 가중치 2.0
            w = 2.0 if i == 0 else 1.0
            s += w * score_text_on_page(ss, ptxt)
        scores.append(s)
    return scores

def score_paragraph_on_pages(text: str, pdf_pages: List[str]) -> List[float]:
    return [score_text_on_page(text, ptxt) for ptxt in pdf_pages]

# -----------------------
# Page assignment
# -----------------------

def assign_pages_to_items(
    items: List[DocxItem],
    pdf_pages: List[str],
) -> Dict[int, List[Dict[str, Any]]]:
    """
    DOCX 아이템을 PDF 페이지에 배치.
    - 순서 보존 (current_page 이상만 선택)
    - 표는 가중치 매칭
    - 같은 시작 셀(혹은 같은 시그니처)을 가진 표가 연속으로 나오면 다른 페이지로 강제 분산
    """
    total_pages = len(pdf_pages)
    page_map: Dict[int, List[Dict[str, Any]]] = {i: [] for i in range(1, total_pages + 1)}

    current_page = 1
    last_table_sig: Optional[str] = None
    last_table_assigned_page: Optional[int] = None

    for kind, payload in items:
        # 1) 점수 벡터
        if kind == "paragraph":
            text = payload.text
            scores = score_paragraph_on_pages(text, pdf_pages)
        else:
            scores = score_table_on_pages(payload, pdf_pages)

        # 2) current_page 이상에서 최고 점수 페이지 선택
        best_page = current_page
        best_score = -1.0
        for idx in range(current_page - 1, total_pages):
            sc = scores[idx]
            if sc > best_score:
                best_score = sc
                best_page = idx + 1

        # 3) 표 병합 방지(연속 동일 표 시그니처면 다음 페이지로 밀어줌)
        if kind == "table":
            # 시그니처 = 헤더 후보 + 첫 본문 문구 몇 개
            sig = "|".join(sample_strings_from_table(payload, max_samples=3))[:128]
            if last_table_sig and sig and normalize(last_table_sig) == normalize(sig):
                if last_table_assigned_page is not None and best_page <= last_table_assigned_page:
                    # 동일 시그니처 표가 같은/이전 페이지로 가려 하면 한 페이지 뒤로
                    best_page = min(total_pages, (last_table_assigned_page + 1))
            last_table_sig = sig
            last_table_assigned_page = best_page

        # 4) 아이템 → JSON 블록 변환
        if kind == "paragraph":
            blk = {"label": text, "content": []} if is_label(payload.text) \
                else {"sen": payload.text}
        else:
            blk = {"table": payload.cells}

        # 5) 페이지에 삽입 및 current_page 갱신 규칙
        page_map[best_page].append(blk)
        # 문단/표가 미래 페이지로 갔으면 current_page 를 따라감(역행 금지)
        if best_page > current_page:
            current_page = best_page

    return page_map

# -----------------------
# Public: build_pages
# -----------------------

def build_pages(docx_path: str, pdf_path: str = "", **kwargs) -> Dict[str, Any]:
    """
    checkreport_api.py 가 호출하는 엔트리 포인트.
    사용법: build_pages(docx_path, pdf_path=pdf_path)  # 키워드 인자 허용
    반환 JSON 스키마: v=0.4 / total_pages / pages[ {page, header, footer, content} ]
    """
    if not pdf_path:
        raise ValueError("pdf_path 가 필요합니다. build_pages(docx_path, pdf_path=...) 형태로 호출하세요.")

    pdf_pages_text = extract_pdf_pages_text(pdf_path)
    items = extract_docx_items(docx_path)
    page_map = assign_pages_to_items(items, pdf_pages_text)

    total_pages = len(pdf_pages_text)

    # header/footer: 간단 버전 — 페이지 텍스트 상위/하위 몇 토큰에서 반복 패턴을 골라낼 수도 있지만
    # 안정성을 위해 비워두거나 필요한 경우 나중에 채워넣도록 둔다.
    pages_json: List[Dict[str, Any]] = []
    for pno in range(1, total_pages + 1):
        pages_json.append({
            "page": pno,
            "header": [],   # 필요하다면 후속 단계에서 주입
            "footer": [],
            "content": page_map.get(pno, []),
        })

    return {
        "v": "0.4",
        "total_pages": total_pages,
        "pages": pages_json,
    }

# -----------------------
# Optional: CLI (debug)
# -----------------------

if __name__ == "__main__":
    import sys
    if len(sys.argv) < 3:
        print("Usage: python report_docx_parser.py <docx_path> <pdf_path>")
        sys.exit(1)
    docx_p, pdf_p = sys.argv[1], sys.argv[2]
    data = build_pages(docx_p, pdf_path=pdf_p)
    print(json.dumps(data, ensure_ascii=False, indent=2))
